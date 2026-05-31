#!/usr/bin/env python3
"""
Generate a highly professional, well-formatted Word (.docx) document
and a plain text file (.txt) containing the Research Highlights
for the CEUS submission, strictly complying with the Elsevier character limit
(3 to 5 bullet points, maximum 85 characters per bullet point, including spaces).
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    highlights = [
        "An open-source Python framework integrates GNNs, RAG-LLMs, and multi-agent systems.",
        "Models urban parcels as a heterogeneous graph with five distinct edge types.",
        "A RAG-LLM parser digitizes complex zoning codes with over 96% extraction accuracy.",
        "MARL agents negotiate multi-objective land-use and stakeholder trade-offs.",
        "Validated on Manhattan's 42,075 parcels, satisfying infrastructure limits."
    ]

    # Print validation of character counts
    print("Validating character counts (Elsevier limit: <= 85 characters including spaces):")
    for idx, highlight in enumerate(highlights, 1):
        char_count = len(highlight)
        print(f"  Bullet {idx}: '{highlight}' -> {char_count} characters")
        assert char_count <= 85, f"Bullet {idx} exceeds 85 characters!"

    # 1. Generate Word Document (.docx)
    doc = Document()
    
    # 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Header
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_header.paragraph_format.space_before = Pt(12)
    p_header.paragraph_format.space_after = Pt(24)
    run_header = p_header.add_run("Research Highlights\nComputers, Environment and Urban Systems (CEUS)")
    run_header.font.name = 'Calibri'
    run_header.font.size = Pt(14)
    run_header.bold = True
    run_header.font.color.rgb = RGBColor(31, 78, 121)

    # Context / Intro paragraph
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(18)
    p_intro.paragraph_format.line_spacing = 1.15
    run_intro = p_intro.add_run(
        "Title of Manuscript: PIMALUOS: An Open-Source Physics-Informed Multi-Agent Framework for Urban Land-Use Optimization\n"
        "Corresponding Author: Parya Payami (ppayami@uwm.edu)\n\n"
        "Highlights consist of 3 to 5 bullet points that convey the core findings and provide a quick textual overview of the article. "
        "Each bullet point is limited to a maximum of 85 characters, including spaces, in accordance with Elsevier guidelines:"
    )
    run_intro.font.name = 'Calibri'
    run_intro.font.size = Pt(11)
    run_intro.italic = True

    # Bullet points
    for highlight in highlights:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.paragraph_format.space_after = Pt(8)
        p_bullet.paragraph_format.line_spacing = 1.15
        run_bullet = p_bullet.add_run(highlight)
        run_bullet.font.name = 'Calibri'
        run_bullet.font.size = Pt(11)

    docx_path = Path("./Highlights.docx")
    doc.save(docx_path)
    print(f"Successfully saved Highlights.docx to {docx_path.resolve()}")

    # 2. Generate Text Document (.txt)
    txt_path = Path("./Highlights.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("PIMALUOS - Research Highlights (CEUS Submission)\n")
        f.write("=================================================\n\n")
        f.write("Elsevier Guidelines Checklist:\n")
        f.write("- 3 to 5 bullet points\n")
        f.write("- Maximum of 85 characters per bullet point (including spaces)\n")
        f.write("- Written in the present tense, stating main findings/methods\n\n")
        f.write("Research Highlights:\n")
        for highlight in highlights:
            f.write(f"- {highlight} ({len(highlight)} chars)\n")
            
    print(f"Successfully saved Highlights.txt to {txt_path.resolve()}")

if __name__ == "__main__":
    main()
