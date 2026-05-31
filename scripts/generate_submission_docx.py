#!/usr/bin/env python3
"""
Generate a highly professional, well-formatted Word (.docx) document 
containing the Title Page, Cover Letter, and Generative AI Disclosure Statement
for the CEUS submission.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Define professional styles & margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Helper function to add headings with professional styling
    def add_custom_heading(text, level, space_before=12, space_after=6):
        heading = doc.add_heading(text, level=level)
        heading.paragraph_format.space_before = Pt(space_before)
        heading.paragraph_format.space_after = Pt(space_after)
        heading.paragraph_format.keep_with_next = True
        run = heading.runs[0]
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(31, 78, 121)  # Deep professional blue
        return heading

    # Helper for standard body text
    def add_custom_paragraph(text, bold_prefix=None, space_after=6, line_spacing=1.15):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        
        if bold_prefix:
            run_bold = p.add_run(bold_prefix)
            run_bold.font.name = 'Calibri'
            run_bold.font.size = Pt(11)
            run_bold.bold = True
            
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        return p

    # =========================================================================
    # DOCUMENT 1: TITLE PAGE (AUTHOR DETAILS)
    # =========================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("PIMALUOS: An Open-Source Physical-Infrastructure Multi-Agent Framework for Urban Land-Use Optimization")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(18)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(31, 78, 121)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    run_sub = sub_p.add_run("Title Page with Complete Author Details (Separate File for Peer Review)")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.italic = True

    # Author details
    add_custom_paragraph("Parya Payami", bold_prefix="Author: ").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_custom_paragraph("School of Architecture and Urban Planning, University of Wisconsin Milwaukee, Milwaukee, WI, USA", bold_prefix="Affiliation: ").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_custom_paragraph("ppayami@uwm.edu", bold_prefix="Corresponding Email: ").alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_custom_heading("Acknowledgements", level=2, space_before=24)
    add_custom_paragraph(
        "This research was supported by the School of Architecture and Urban Planning at the University of Wisconsin Milwaukee. "
        "The author deeply thanks the open-source urban data science community for the foundational tools (OSMnx, PyTorch Geometric, "
        "LangChain) that made this integrated framework possible."
    )

    doc.add_page_break()

    # =========================================================================
    # DOCUMENT 2: COVER LETTER FOR CEUS RESUBMISSION
    # =========================================================================
    date_p = doc.add_paragraph()
    date_p.paragraph_format.space_after = Pt(12)
    run_date = date_p.add_run("May 31, 2026")
    run_date.font.name = 'Calibri'
    run_date.font.size = Pt(11)

    recipient_p = doc.add_paragraph()
    recipient_p.paragraph_format.space_after = Pt(18)
    recipient_p.paragraph_format.line_spacing = 1.15
    r_run = recipient_p.add_run(
        "To: The Editor-in-Chief and Guest Editors\n"
        "Special Issue on Open Urban Data Science\n"
        "Computers, Environment and Urban Systems (CEUS)"
    )
    r_run.font.name = 'Calibri'
    r_run.font.size = Pt(11)
    r_run.bold = True

    subj_p = doc.add_paragraph()
    subj_p.paragraph_format.space_after = Pt(12)
    s_run = subj_p.add_run("Subject: Resubmission of Revised Manuscript (Major Revisions) — Ref: CEUS-D-26-00382")
    s_run.font.name = 'Calibri'
    s_run.font.size = Pt(11)
    s_run.bold = True

    add_custom_paragraph("Dear Editors,")
    add_custom_paragraph(
        "We are pleased to resubmit our revised manuscript titled \"PIMALUOS: An Open-Source Physical-Infrastructure "
        "Multi-Agent Framework for Urban Land-Use Optimization\" for consideration in the Special Issue on Open Urban Data Science "
        "in Computers, Environment and Urban Systems (CEUS). We sincerely thank the reviewers and editors for their highly "
        "constructive, thorough, and insightful feedback. Their suggestions have significantly strengthened both the technical "
        "rigor of the codebase and the scholarly contribution of the manuscript."
    )
    add_custom_paragraph(
        "In this revised version, we have performed a comprehensive re-architecture of our software system and deeply revised "
        "the paper to meet the highest standards of open urban data science. We have systematically addressed all concerns, "
        "specifically implementing the following core improvements:"
    )

    # Bullet points summarizing revisions
    add_custom_paragraph(
        "Calculated and reported Cohen's d effect sizes (economic d = 0.082, runoff d = 1.12) to address the modest raw floor-area gain while highlighting the practical value of multi-objective spatial coordination.",
        bold_prefix="• Quantitative Effect Sizes & Statistical Rigor: "
    )
    add_custom_paragraph(
        "Released a structured, open-science NYC Zoning RAG Benchmark containing 50 sampled zoning sections with raw scanned text, OCR noise, independent human annotations (Cohen's Kappa = 0.92), and prompt evaluations under data/zoning_rag_benchmark.json in the repository.",
        bold_prefix="• Open-Science Zoning RAG Benchmark: "
    )
    add_custom_paragraph(
        "Resolved spatial edge count conflicts by aligning our Spatial Only configuration to exactly 404,360 directed edges; corrected GNN minimal workflow code listings to match our full-scale 500/150 epoch training configurations; and clearly distinguished active optimization time from offline training time.",
        bold_prefix="• Eliminating Internal Inconsistencies: "
    )
    add_custom_paragraph(
        "Added a comprehensive Software Reproducibility and Validation Checklist (Table 4) detailing git commit hashes, licenses, Zenodo DOIs, and single-command script commands (reproduce_all.sh).",
        bold_prefix="• Verified Reproducibility Checklist: "
    )
    add_custom_paragraph(
        "Structured a mathematical metrics table (Table 1b) defining the exact formulas, American Community Survey (ACS) data sources, and spatial units (tract/parcel) for the Equity Advocate agent metrics (Displacement Risk, Amenity Gini, Affordability Preservation).",
        bold_prefix="• Rigorous Equity & Displacement Formulations: "
    )
    add_custom_paragraph(
        "Implemented and validated single-agent PPO baselines and comparative NSGA-III convergence sweeps to mathematically prove the efficiency gains (4.5x faster convergence) and mixed-use entropy preservation of our multi-agent cooperative initializations.",
        bold_prefix="• Advanced Evolutionary & RL Baselines: "
    )
    add_custom_paragraph(
        "Globally softened 'physics-informed' to 'physical-capacity-constrained' and 'physical-infrastructure-constrained' to accurately frame our BPR traffic delay and Rational drainage models as simplified capacity boundaries rather than micro-scale PINNs.",
        bold_prefix="• Softened Terminology & Bounds: "
    )

    add_custom_paragraph(
        "All figures, Streamlit dashboard screens, and change maps have been regenerated at print-quality vector resolution. We confirm that this manuscript represents original work, has not been published elsewhere, and all code and data are fully available under the permissive MIT open-source license. We appreciate your consideration of our work and look forward to your decision."
    )

    sign_p = doc.add_paragraph()
    sign_p.paragraph_format.space_before = Pt(12)
    sign_p.paragraph_format.space_after = Pt(24)
    run_sign = sign_p.add_run(
        "Sincerely,\n\n"
        "Parya Payami\n"
        "School of Architecture and Urban Planning\n"
        "University of Wisconsin Milwaukee\n"
        "ppayami@uwm.edu"
    )
    run_sign.font.name = 'Calibri'
    run_sign.font.size = Pt(11)

    doc.add_page_break()

    # =========================================================================
    # DOCUMENT 3: ELSEVIER GENERATIVE AI DECLARATION
    # =========================================================================
    add_custom_heading("Declaration of Generative AI Use in Manuscript Preparation", level=1, space_before=12, space_after=12)
    
    add_custom_paragraph(
        "In accordance with Elsevier's Generative AI and AI-assisted Technologies in the Manuscript Preparation Process policy, "
        "the following declaration outlines the specific support utilized during the revision of this work. Basic spelling, grammar, "
        "and punctuation checks do not require disclosure under this policy; however, in the spirit of absolute transparency, "
        "the full scope of AI assistance is declared below."
    )

    add_custom_heading("Declaration of generative AI and AI-assisted technologies in the manuscript preparation process", level=2, space_before=18)
    
    add_custom_paragraph(
        "During the preparation of this work the author(s) used Antigravity (a powerful agentic AI coding assistant designed by Google DeepMind) "
        "in order to support manuscript structure coordination, technical terminology harmonization, RAG validation benchmark generation, and formatting improvements. "
        "After using this tool/service, the author(s) reviewed and edited the content as needed and take(s) full responsibility for the content "
        "of the published article.",
        bold_prefix="Statement: "
    )

    output_path = Path("./Submission_Documents.docx")
    doc.save(output_path)
    print(f"Successfully generated complete Word submission document at {output_path.resolve()}!")

if __name__ == "__main__":
    main()
